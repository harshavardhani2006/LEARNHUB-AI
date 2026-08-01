import io

class DocumentParser:
    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF file bytes using multiple strategies."""
        text = ""

        # Strategy 1: pypdf (newer, better layout handling)
        try:
            from pypdf import PdfReader as NewPdfReader
            pdf_file = io.BytesIO(file_bytes)
            reader = NewPdfReader(pdf_file)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
            if text.strip():
                return text
        except Exception:
            pass

        # Strategy 2: PyPDF2 fallback
        try:
            import PyPDF2
            pdf_file = io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text() or ""
                text += page_text + "\n"
            if text.strip():
                return text
        except Exception:
            pass

        # Strategy 3: If no text extracted (image-based PDF),
        # return a meaningful placeholder so study tools don't crash
        if not text.strip():
            return (
                "This document appears to be image-based or has complex formatting. "
                "Text extraction was limited. The AI tools will provide general guidance "
                "based on the document title and subject category."
            )

        return text

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX file bytes."""
        text = ""
        try:
            from docx import Document
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
        return text

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        """Extract text from TXT file bytes."""
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            raise ValueError(f"Failed to parse TXT: {str(e)}")

    @classmethod
    def parse(cls, filename: str, file_bytes: bytes) -> str:
        """Parse document based on its extension."""
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return cls.parse_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            return cls.parse_docx(file_bytes)
        elif ext == "txt":
            return cls.parse_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: .{ext}")
